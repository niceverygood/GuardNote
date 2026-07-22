from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "GuardNote_사용설명서.docx"
OG_IMAGE = ROOT / "public" / "og.png"

# compact_reference_guide preset + named GuardNote brand overrides.
# Use a system-wide Unicode font so LibreOffice can embed Hangul when it runs
# under an isolated HOME during automated rendering.
FONT = "Arial Unicode MS"
INK = RGBColor(0x20, 0x1F, 0x1D)
MUTED = RGBColor(0x72, 0x6D, 0x66)
ORANGE = RGBColor(0xEB, 0x69, 0x17)
ORANGE_DARK = RGBColor(0xC9, 0x4F, 0x06)
PEACH = "FFF1E7"
CREAM = "F7F5F1"
GREEN = RGBColor(0x23, 0x84, 0x5A)
GREEN_FILL = "E8F5EF"
RED = RGBColor(0xBC, 0x3E, 0x3E)
RED_FILL = "FAEAEA"
LINE = "E7E3DC"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 55
CELL_SIDE = 120


def set_run(run, size=10.25, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_SIDE, bottom=CELL_TOP_BOTTOM, end=CELL_SIDE):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    max_abs = max([int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))] or [0])
    ids = {}
    for kind, fmt, text, font in (("decimal", "decimal", "%1.", FONT), ("bullet", "bullet", "●", FONT)):
        max_abs += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(max_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        fonts.set(qn("w:eastAsia"), font)
        fonts.set(qn("w:cs"), font)
        fonts.set(qn("w:hint"), "eastAsia")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        max_num += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(max_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(max_abs))
        num.append(abs_id)
        numbering.append(num)
        ids[kind] = max_num
    return ids


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.76)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.8)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, before, after in (("Heading 1", 15, 14, 8), ("Heading 2", 12, 10, 5), ("Heading 3", 11, 8, 4)):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ORANGE_DARK if name != "Heading 3" else INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    # The builder styles headers/footers both before and after composing pages.
    # Clear previously inserted runs/fields so Word does not duplicate the footer.
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("GuardNote 사용자 가이드   |   ")
    set_run(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.text = "GUARD NOTE   /   PRIVACY INTELLIGENCE"
    set_run(hp.runs[0], size=8, bold=True, color=MUTED)
    hp.paragraph_format.space_after = Pt(0)
    add_page_field(section.footer.paragraphs[0])


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text.upper()), size=8, bold=True, color=ORANGE)
    return p


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_run(p.add_run(text), size=25, bold=True, color=INK)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(12)
        set_run(sp.add_run(subtitle), size=12, color=MUTED)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=INK)
        set_run(p.add_run(text[len(bold_lead):]), color=INK)
    else:
        set_run(p.add_run(text), color=INK)
    return p


def add_bullet(doc, text, nums, bold_lead=None):
    p = doc.add_paragraph()
    set_numbering(p, nums["bullet"])
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_step(doc, text, nums):
    p = doc.add_paragraph()
    set_numbering(p, nums["decimal"])
    set_run(p.add_run(text))
    return p


def add_callout(doc, title, text, tone="orange"):
    fills = {"orange": PEACH, "green": GREEN_FILL, "red": RED_FILL, "gray": CREAM}
    colors = {"orange": ORANGE_DARK, "green": GREEN, "red": RED, "gray": INK}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), size=10, bold=True, color=colors[tone])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(text), size=9, color=INK)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 3360, 4100])
    headers = ["메뉴", "무엇을 하는 곳인가", "핵심 결과"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PEACH)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(text), size=8.25, bold=True, color=ORANGE_DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(text), size=8.25)
        set_table_geometry(table, [1900, 3360, 4100])
    return table


def add_section_break(doc, title, subtitle, number):
    doc.add_page_break()
    add_kicker(doc, f"{number:02d} / USER GUIDE")
    add_title(doc, title, subtitle)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    for section in doc.sections:
        set_header_footer(section)
    nums = configure_numbering(doc)

    # Editorial cover.
    cover = doc.sections[0]
    cover.different_first_page_header_footer = True
    doc.add_paragraph().paragraph_format.space_after = Pt(34)
    add_kicker(doc, "GUARDNOTE PRODUCT MANUAL · 2026.07")
    add_title(doc, "GuardNote 사용설명서", "개인정보 준수 업무와 AI 문서 자동화를 위한 실무 가이드")
    add_body(doc, "자가점검, 수탁자 관리, Privacy by Design, AI 문서검토·자동답변·문서생성·증적검토를 한 흐름으로 사용하는 방법을 안내합니다.")
    if OG_IMAGE.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(OG_IMAGE), width=Inches(6.3))
    add_callout(doc, "버전 정보", "기준일 2026.07.22 · 개인정보의 안전성 확보조치 기준 2025.11 반영 · 로컬 실행 주소 http://127.0.0.1:5173/", "gray")

    add_section_break(doc, "3분 빠른 시작", "처음 접속한 사용자가 가장 빠르게 결과를 확인하는 순서", 1)
    for step in (
        "왼쪽 메뉴에서 개인정보 자가점검을 열고 현재 답변 진행률과 위반 가능 건수를 확인합니다.",
        "AI 자동답변을 선택하고 샘플 동의서 또는 검토할 문서를 첨부합니다.",
        "AI가 제안한 답변, 정확도, 문서 근거를 확인한 후 필요한 답변만 점검표에 반영합니다.",
        "수탁자 점검에서 증적자료를 연결하고 AI 증적 검토로 답변과의 일치 여부를 판정합니다.",
        "Smart Docs에서 문서검토 또는 문서생성을 실행하고 결과를 다운로드합니다.",
    ):
        add_step(doc, step, nums)
    add_callout(doc, "권장 순서", "자가점검 → 수탁자 점검 → PbD 프로젝트 → Smart Docs → AI 검토 내역 순으로 보면 전체 준수 흐름을 가장 빠르게 이해할 수 있습니다.", "green")
    doc.add_heading("화면 구성", level=2)
    add_feature_table(doc, [
        ("운영 대시보드", "준수 상태와 우선 조치 업무 요약", "점수, 진행률, 열린 위험, 최근 AI 작업"),
        ("개인정보 자가점검", "약 600개 기준 기반 즉시 진단", "답변, 위반·권고, 법적 근거, 보고서"),
        ("수탁자 점검", "답변·증적·교육을 연결한 관리감독", "응답률, 점수, AI 증적 판정, 보완 요청"),
        ("PbD 프로젝트", "제품·서비스 출시 전 위험 관리", "프로젝트별 점수, 게이트, 열린 위험"),
        ("Smart Docs", "문서의 검토·답변·생성·증적 분석", "위반·권고, 자동답변, 편집 가능한 초안"),
        ("AI 검토 내역", "AI 결과와 담당자 결정 추적", "신뢰도, 상태, 처리 시간, CSV"),
    ])

    add_section_break(doc, "운영 대시보드", "위험과 업무 우선순위를 먼저 보는 시작 화면", 2)
    doc.add_heading("주요 지표 읽기", level=2)
    for item in (
        "현재 준수 점수: 자가점검, 수탁자, 프로젝트의 주요 위험을 요약한 운영 지표입니다.",
        "점검 완료: 전체 600개 기준 중 현재 답변이 완료된 수를 표시합니다.",
        "수탁자 응답률: 점검을 요청한 수탁자 가운데 답변을 완료한 비율입니다.",
        "열린 위험: 담당자의 조치나 추가 증적이 필요한 위반·권고 건수입니다.",
        "AI 자동화: 이번 달 문서검토, 자동답변, 문서생성, 증적검토 실행 건수입니다.",
    ):
        add_bullet(doc, item, nums, item.split(":")[0] + ":")
    doc.add_heading("오늘 먼저 처리할 일", level=2)
    add_body(doc, "긴급도와 기한을 기준으로 조치가 필요한 업무를 위쪽에 배치합니다. 항목을 선택하면 해당 수탁자, 문서 또는 프로젝트 화면으로 이동합니다.")
    add_callout(doc, "실무 팁", "매일 업무 시작 시 대시보드의 긴급 항목과 최근 AI 검토를 먼저 확인하면 증적 보완과 문서 승인 누락을 줄일 수 있습니다.", "orange")

    add_section_break(doc, "개인정보 자가점검", "답변 즉시 위반 가능성과 개선 권고를 확인하는 방법", 3)
    doc.add_heading("수동 답변", level=2)
    for step in (
        "질문 영역 필터에서 전체, 동의, 제3자 제공, 아동, 국외 이전 중 필요한 범위를 선택합니다.",
        "각 질문에서 예, 아니요, 해당 없음 중 하나를 선택합니다.",
        "오른쪽 판단 가이드에서 관계 법령, 판단 기준, 권장 증적을 확인합니다.",
        "필요하면 증적자료 연결을 선택해 화면 캡처, 승인 이력, 최신 문서를 연결합니다.",
        "보고서를 선택해 현재 답변과 법적 근거를 텍스트 파일로 저장합니다.",
    ):
        add_step(doc, step, nums)
    doc.add_heading("답변 상태 해석", level=2)
    add_feature_table(doc, [
        ("예", "요구되는 통제나 절차가 확인된 상태", "준수 상태로 표시"),
        ("아니요", "필수 절차 또는 문서 근거가 부족한 상태", "위반 가능 또는 개선 필요"),
        ("해당 없음", "현재 사업·처리 흐름에 적용되지 않는 상태", "적용 제외 사유 확인 필요"),
    ])
    add_callout(doc, "주의", "‘예’ 답변만으로 준수가 확정되는 것은 아닙니다. 실제 화면, 문서, 승인 기록 등 객관적인 증적을 함께 연결해야 합니다.", "red")

    add_section_break(doc, "AI 자동답변", "동의서와 처리방침에서 질문별 답변과 근거를 찾는 방법", 4)
    for step in (
        "개인정보 자가점검 또는 Smart Docs에서 AI 자동답변을 엽니다.",
        "PDF, DOCX, HWP, TXT 문서를 선택하거나 샘플 동의서를 사용합니다.",
        "분석 범위를 확인하고 AI 답변 생성을 실행합니다.",
        "문항별 AI 답변, 정확도, 문서 근거를 읽습니다.",
        "근거가 충분한 항목만 답변 반영을 선택합니다.",
    ):
        add_step(doc, step, nums)
    doc.add_heading("결과 화면에서 확인할 것", level=2)
    for item in (
        "AI 답변: 예, 아니요, 해당 없음 중 제안된 값",
        "정확도: 해당 문서가 질문의 근거로 충분하다고 판단한 신뢰도",
        "문서 근거: 답변을 제안한 문구 또는 누락된 정보",
        "반영 결과: 선택한 답변이 자가점검 진행률과 위반 건수에 즉시 반영되는지 여부",
    ):
        add_bullet(doc, item, nums, item.split(":")[0] + ":")
    add_callout(doc, "개인정보 보호", "현재 데모 분석은 브라우저 안에서 실행되며 선택한 파일을 외부 서비스로 전송하지 않습니다. 운영용 AI를 연결할 때에도 API 키는 서버 환경변수로만 관리해야 합니다.", "green")

    add_section_break(doc, "수탁자 점검과 교육", "다수 수탁자의 답변, 증적, 교육 상태를 한 번에 관리하는 방법", 5)
    doc.add_heading("수탁자 등록과 점검", level=2)
    for step in (
        "수탁자 추가를 선택하고 회사명과 위탁업무를 입력합니다.",
        "필수 점검표의 관리적, 기술적, 개인정보 생명주기 문항 구성을 확인합니다.",
        "수탁자별 진행률, 점수, 증적 건수, 보완 상태를 비교합니다.",
        "교육 미이수 수탁자에게 교육을 요청하고, 이수한 경우 수료증을 저장합니다.",
    ):
        add_step(doc, step, nums)
    doc.add_heading("관리자가 볼 지표", level=2)
    add_feature_table(doc, [
        ("진행률", "수탁자가 답변한 점검 문항 비율", "미응답 독촉과 마감 관리"),
        ("점수", "답변과 증적을 반영한 준수 수준", "보완 우선순위 설정"),
        ("증적", "수탁자가 제출한 문서·화면 수", "객관적인 이행 여부 검증"),
        ("교육", "개인정보 처리업무 교육 이수 여부", "수료증 발급과 이수 추적"),
    ])

    add_section_break(doc, "AI 증적검토", "수탁자의 답변과 제출 자료가 실제로 일치하는지 판정하는 방법", 6)
    for step in (
        "수탁자 목록에서 검토할 회사를 선택합니다.",
        "검토 문항의 현재 답변을 확인합니다.",
        "매뉴얼, 화면 캡처, 교육·훈련 결과 등 증적자료를 선택합니다.",
        "AI 증적 검토를 실행합니다.",
        "일치 또는 불일치 판정, 신뢰도, 판정 근거, 보완 제안을 확인합니다.",
        "불일치인 경우 보완 요청 생성을 선택해 담당자에게 필요한 자료를 명확하게 안내합니다.",
    ):
        add_step(doc, step, nums)
    add_callout(doc, "불일치 예시", "‘침해사고 대응 절차’ 질문에 단순한 ID 찾기 화면을 제출하면 담당자 역할, 보고 체계, 정보주체 통지, 관계기관 신고 절차를 확인할 수 없어 불일치로 판정됩니다.", "red")
    add_callout(doc, "권장 보완자료", "침해사고 대응 매뉴얼, 비상연락망, 최근 모의훈련 결과, 사고 보고·승인 이력 중 질문과 직접 연결되는 자료를 요청하세요.", "orange")

    add_section_break(doc, "PbD 프로젝트", "새 제품과 서비스의 설계 단계에서 위험을 없애는 방법", 7)
    doc.add_heading("프로젝트별 관리", level=2)
    for item in (
        "프로젝트 카드에서 설계 검토, 개선 중, 출시 승인 등 현재 단계를 확인합니다.",
        "프로젝트별 점수와 열린 위험 수를 비교해 검토 우선순위를 정합니다.",
        "기획, 설계 검토, 개발, 출시 승인 게이트를 따라 체크리스트를 완료합니다.",
        "체크 항목을 확인하면 조치 필요 수가 실시간으로 줄어들고 모두 완료하면 검토 완료로 표시됩니다.",
        "보고서를 내려받아 출시 승인 기록과 개인정보 영향 검토 자료로 활용합니다.",
    ):
        add_bullet(doc, item, nums)
    doc.add_heading("필수 설계 점검", level=2)
    for item in (
        "처리 목적과 최소 수집항목",
        "프로파일링과 자동화된 결정 안내",
        "보유기간과 파기 트리거",
        "관리자 최소권한과 접속기록",
        "수탁자·재수탁자와 국외 이전 흐름",
        "출시 전 개인정보 영향 검토와 승인 책임자",
    ):
        add_bullet(doc, item, nums)

    add_section_break(doc, "AI 문서검토", "개인정보 문서의 위반, 권고, 개선안과 근거를 확인하는 방법", 8)
    for step in (
        "Smart Docs에서 AI 문서 검토를 선택합니다.",
        "동의서, 개인정보 처리방침, 내부 관리계획, 위수탁계약서 중 문서 유형을 선택합니다.",
        "검토할 문서를 첨부하고 아동, 민감정보, 마케팅, 국외 이전 등 적용 맥락을 지정합니다.",
        "AI 정확성 검토를 실행합니다.",
        "문서 점수와 위반·권고·준수 건수를 확인합니다.",
        "각 결과의 문제 문구, 개선 방향, 관계 법령을 검토한 후 보고서를 저장합니다.",
    ):
        add_step(doc, step, nums)
    doc.add_heading("심각도 구분", level=2)
    add_feature_table(doc, [
        ("위반", "필수 고지 또는 동의 요건이 누락·왜곡된 상태", "수정 후 승인 필요"),
        ("권고", "법적 근거 또는 표현이 모호해 추가 확인이 필요한 상태", "담당자 검토와 개선 권장"),
        ("준수", "현재 문구에서 필수 요건과 근거가 확인된 상태", "최종 승인 전 문맥 확인"),
    ])
    add_callout(doc, "대표 검토 항목", "선택 동의 분리, 제3자 제공 보유기간, 고유식별정보 처리 근거, 제3자 연락처, 중요사항 강조 표시를 중점적으로 확인합니다.", "orange")

    add_section_break(doc, "AI 문서생성", "관련 정보로 문서 초안을 만들고 편집·검증·다운로드하는 방법", 9)
    for step in (
        "Smart Docs에서 AI 문서 생성을 선택합니다.",
        "동의서, 처리방침, 내부 관리계획, 위수탁계약서 중 만들 문서를 선택합니다.",
        "회사·기관명, 처리 목적, 개인정보 항목, 보유기간, 문의처를 입력합니다.",
        "안전한 초안 생성을 실행합니다.",
        "필수 항목 확인 결과와 담당자 검토 경고를 확인합니다.",
        "생성된 초안을 직접 편집하고 조항별 법적 근거를 확인합니다.",
        "다운로드를 선택해 텍스트 초안을 저장합니다.",
    ):
        add_step(doc, step, nums)
    add_callout(doc, "출시 예정 기능을 완성한 부분", "문서 유형 선택뿐 아니라 필수 정보 입력, 법적 조항 구성, 누락 경고, 편집 가능한 결과, 조항 근거, 다운로드까지 하나의 실제 업무 흐름으로 구현했습니다.", "green")
    add_callout(doc, "최종 승인 원칙", "생성 문서는 법률 자문이나 확정 문서가 아닙니다. 실제 처리 흐름, 시스템 설정, 보유기간, 수탁자와 국외 이전 정보를 담당자가 대조한 후 승인하세요.", "red")

    add_section_break(doc, "AI 검토 내역과 보고서", "자동화 결과를 추적하고 내부 감사 자료로 활용하는 방법", 10)
    doc.add_heading("AI 검토 내역", level=2)
    for item in (
        "문서 검토, 자동 답변, 증적 검토, 문서 생성 유형으로 필터링할 수 있습니다.",
        "검토 대상, AI 결과, 신뢰도, 실행 시간, 담당자 처리 상태를 한 줄에서 확인합니다.",
        "항목을 선택하면 설명 가능한 근거와 처리 상세를 확인할 수 있습니다.",
        "CSV 내보내기로 전체 이력을 감사·보고용 자료로 저장합니다.",
    ):
        add_bullet(doc, item, nums)
    doc.add_heading("다운로드되는 파일", level=2)
    add_feature_table(doc, [
        ("자가점검", "GuardNote_개인정보_자가점검_보고서.txt", "질문, 답변, 관계 법령"),
        ("PbD", "[프로젝트명]_PbD_검토보고서.txt", "체크 항목과 조치 상태"),
        ("문서검토", "AI_문서검토_결과.txt", "위반·권고, 개선안, 근거"),
        ("문서생성", "GuardNote_[문서유형]_초안.txt", "편집한 문서 초안"),
        ("검토내역", "GuardNote_AI검토내역.csv", "AI 작업 이력과 상태"),
    ])

    add_section_break(doc, "운영 원칙과 문제 해결", "개인정보와 AI를 안전하게 사용하기 위한 최종 체크리스트", 11)
    doc.add_heading("AI 운영 원칙", level=2)
    for item in (
        "AI 결과는 법률 자문을 대체하지 않으며 담당자가 최종 확정합니다.",
        "신뢰도가 높더라도 문서의 실제 적용 맥락과 시스템 설정을 함께 확인합니다.",
        "민감한 실제 파일을 외부 AI에 보내기 전 계약, 보관, 재학습, 국외 이전 조건을 검토합니다.",
        "API 키는 화면 코드나 문서에 저장하지 않고 서버 환경변수 또는 비밀관리 서비스로 관리합니다.",
        "AI 판정을 수정한 경우 수정 이유와 담당자, 시간을 검토 이력에 남깁니다.",
    ):
        add_bullet(doc, item, nums)
    doc.add_heading("문제가 있을 때", level=2)
    add_feature_table(doc, [
        ("버튼 반응 없음", "페이지를 새로고침하고 다시 실행", "입력한 데모 상태는 초기화될 수 있음"),
        ("분석이 시작되지 않음", "문서 또는 증적 선택 여부 확인", "샘플 문서로 기능을 먼저 검증"),
        ("다운로드가 안 됨", "브라우저 다운로드 허용 확인", "팝업·다운로드 차단 설정 점검"),
        ("AI 결과가 부정확", "문서 유형과 적용 맥락 재확인", "담당자 판단으로 수정 후 근거 기록"),
        ("모바일 메뉴 안 보임", "왼쪽 위 메뉴 버튼 선택", "메뉴 이동 후 자동으로 닫힘"),
    ])
    add_callout(doc, "로컬 실행", "프로젝트 폴더에서 개발 서버가 실행 중일 때 http://127.0.0.1:5173/ 로 접속합니다. 화면이 열리지 않으면 서버 상태를 먼저 확인하세요.", "gray")
    add_callout(doc, "보안 알림", "대화나 화면에 노출된 API 키는 폐기하고 새 키를 발급하세요. 새 키도 소스 코드, Git, 브라우저 저장소에 직접 기록하면 안 됩니다.", "red")

    for section in doc.sections:
        set_header_footer(section)
    doc.core_properties.title = "GuardNote 사용설명서"
    doc.core_properties.subject = "개인정보 준수 및 AI 문서 자동화 사용자 가이드"
    doc.core_properties.author = "GuardNote"
    doc.core_properties.keywords = "개인정보, 컴플라이언스, AI 문서검토, 수탁자, Privacy by Design"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
